from __future__ import annotations

import io
import re
from collections import defaultdict

import pdfplumber
import pymupdf
import pymupdf4llm
from docx import Document as _DocxDocument

SECTION_NAMES = {
    # Skills
    "skill": "Skills",
    "skills": "Skills",
    "technical skill": "Skills",
    "technical skills": "Skills",
    "kỹ năng": "Skills",
    "kĩ năng": "Skills",

    # Experience
    "experience": "Experience",
    "work experience": "Experience",
    "professional experience": "Experience",
    "employment history": "Experience",
    "kinh nghiệm": "Experience",
    "kinh nghiệm làm việc": "Experience",
    "kinh nghiệp làm việc": "Experience",

    # Education
    "education": "Education",
    "academic background": "Education",
    "học vấn": "Education",

    # Projects
    "project": "Projects",
    "projects": "Projects",
    "personal project": "Projects",
    "personal projects": "Projects",
    "dự án": "Projects",
    "dự án cá nhân": "Projects",
    "dự án nhóm": "Projects",

    # Certifications
    "certification": "Certifications",
    "certifications": "Certifications",
    "certifications & awards": "Certifications",
    "awards": "Certifications",
    "chứng chỉ": "Certifications",
    "giải thưởng": "Certifications",

    # Languages
    "language": "Languages",
    "languages": "Languages",
    "ngoại ngữ": "Languages",

    # Summary
    "summary": "Summary",
    "professional summary": "Summary",
    "profile": "Summary",
    "objective": "Summary",
    "career objective": "Summary",
    "mục tiêu nghề nghiệp": "Summary",

    # Contact
    "contact": "Contact",
    "personal information": "Contact",
    "thông tin cá nhân": "Contact",

    # Other
    "additional information": "Additional Information",
    "thông tin thêm": "Additional Information",
    "interests": "Interests",
    "hobbies": "Interests",
    "sở thích": "Interests",
}


BULLET_RE = re.compile(
    r"^\s*(?:[-*+•●▪◦‣–—]|(?:\d+)[.)])\s*"
)

MARKDOWN_HEADING_RE = re.compile(
    r"^(#{1,6})\s+(.+?)\s*$"
)

URL_RE = re.compile(
    r"https?://\S+"
    r"|www\.\S+"
    r"|github\.com/\S+|linkedin\.com/\S+|facebook\.com/\S+"
    # Generic scheme-less "domain.tld/path" (catches twitter.com/handle,
    # gitlab.com/..., portfolio-site.dev/... etc. not on the explicit list
    # above).
    r"|\b[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.(?:com|net|io|dev|me|co|vn)/\S+",
    re.IGNORECASE,
)

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)

PHONE_RE = re.compile(
    r"(?<!\d)(?:\+?84[\s.-]?|0)(?:3|5|7|8|9)\d(?:[\s.-]?\d){7,8}(?!\d)"
)

DOB_RE = re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")

MARKDOWN_LINK_RE = re.compile(r"\[[^\]]+\]\([^)]+\)")

LABELED_PII_RE = re.compile(
    r"^\s*(?:[-*+•] )?(?:"
    r"email|e-mail|phone|tel|mobile|điện thoại|dien thoai|"
    r"địa chỉ|dia chi|address|linkedin|github|facebook|twitter|"
    r"ngày sinh|ngay sinh|date of birth|dob|giới tính|gioi tinh|gender|"
    r"cccd|cmnd|passport|họ và tên|ho va ten|full name|name"
    r")\s*[:：].*$",
    re.IGNORECASE,
)

# Below this many characters of clean CV content, extraction is almost
# certainly a parse failure (complex layout, scanned image, corrupted
# file) rather than a genuinely short resume. Tuned against the
# ingest_eval_v2 "hard" real-world PDF sample (worst case: 710 chars).
LOW_CONTENT_CHAR_THRESHOLD = 600

_CONTACT_SECTIONS = {"contact", "personal information"}
_NAME_PARTICLES = {"van", "von", "de", "del", "da", "thi", "la", "le", "binh"}
_JOB_WORDS = {
    "engineer",
    "developer",
    "intern",
    "student",
    "manager",
    "bachelor",
    "master",
    "cv",
    "resume",
}


def _normalize_section_key(text: str) -> str:
    text = text.strip()

    # Existing markdown
    text = re.sub(r"^#{1,6}\s*", "", text)

    # Existing bullets / numbering
    text = BULLET_RE.sub("", text)

    text = text.rstrip(":").strip()
    text = re.sub(r"\s+", " ", text)
    folded = text.casefold()

    # Real-world templates often prefix section titles with a decorative
    # emoji/icon (e.g. "💼 KINH NGHIỆM / DỰ ÁN") — strip any leading
    # non-word characters so the keyword lookup below still matches.
    return re.sub(r"^[^\w]+", "", folded)


# Longest keyword first, so e.g. "kinh nghiệm làm việc" is tried before
# the shorter "kinh nghiệm" prefix.
_SECTION_NAME_ITEMS = sorted(SECTION_NAMES.items(), key=lambda item: len(item[0]), reverse=True)


def _detect_section(text: str) -> tuple[str, str] | None:
    """
    Returns (canonical_section, remainder) if `text` looks like a section
    heading, else None. `remainder` is whatever followed the matched
    keyword (e.g. "chuyên môn" in "KỸ NĂNG CHUYÊN MÔN", or "BEng
    Software, Bach Khoa HCM" in "Education BEng Software, Bach Khoa HCM"
    when a PDF renderer joined a heading and its content onto one line).
    Callers must keep a non-empty remainder as a content line — dropping
    it would silently delete real CV content whenever the remainder
    happens to be actual data rather than a decorative qualifier.
    """
    key = _normalize_section_key(text)
    if key in SECTION_NAMES:
        return SECTION_NAMES[key], ""

    # Fallback: real CVs often pad a section title with extra words
    # ("KỸ NĂNG CHUYÊN MÔN" = "Skills" + "professional", "KINH NGHIỆM /
    # DỰ ÁN" = "Experience" + "/ Projects").
    for name, canonical in _SECTION_NAME_ITEMS:
        if key == name:
            return canonical, ""
        if key.startswith(f"{name} "):
            return canonical, key[len(name) :].strip()

    return None


def _clean_text_artifacts(text: str) -> str:
    """
    Remove common PDF/OCR garbage without aggressively modifying
    legitimate resume content.
    """
    text = text.replace("\x00", "")
    text = text.replace("\ufeff", "")
    text = text.replace("\u00ad", "")

    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove trailing spaces
    text = "\n".join(line.rstrip() for line in text.splitlines())

    # Avoid huge empty gaps
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    return text.strip()


def _looks_corrupted(text: str) -> bool:
    """
    Conservative quality gate.

    Only trigger full OCR when there are clear signs that the native
    PDF text layer is unreliable.
    """
    if not text.strip():
        return True

    replacement_chars = text.count("\ufffd")
    null_chars = text.count("\x00")

    if replacement_chars >= 3 or null_chars >= 1:
        return True

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    if not lines:
        return True

    # Example:
    # https://facebook.com/abc
    # m
    #
    # or highly fragmented extraction.
    tiny_lines = sum(
        1
        for line in lines
        if len(line) <= 2
        and not re.fullmatch(r"\d+", line)
    )

    if len(lines) >= 10 and tiny_lines / len(lines) > 0.10:
        return True

    # Many unicode replacement-like/control characters
    suspicious = sum(
        1
        for ch in text
        if ord(ch) < 32 and ch not in "\n\t"
    )

    if suspicious > 2:
        return True

    return False


def _pdf_to_markdown(
    data: bytes,
    *,
    force_ocr: bool = False,
) -> str:
    """
    Convert PDF bytes directly to layout-aware Markdown.
    """
    with pymupdf.open(
        stream=data,
        filetype="pdf",
    ) as doc:
        markdown = pymupdf4llm.to_markdown(
            doc,
            header=False,
            footer=False,
            use_ocr=True,
            force_ocr=force_ocr,
            ocr_language="vie+eng",
            show_progress=False,
        )

    return _clean_text_artifacts(markdown)


def _normalize_markdown_sections(markdown: str) -> str:
    """
    Standardize known resume sections as ## headings.

    Existing useful Markdown from PyMuPDF4LLM is retained.
    We deliberately avoid converting every short line into ### because
    that heuristic creates many false headings.
    """
    lines = markdown.splitlines()
    output: list[str] = []

    previous_blank = True

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            if output and output[-1] != "":
                output.append("")
            previous_blank = True
            continue

        detected = _detect_section(line)

        if detected:
            section, remainder = detected
            if output and output[-1] != "":
                output.append("")

            output.append(f"## {section}")
            output.append("")
            if remainder:
                output.append(remainder)
            previous_blank = True
            continue

        heading_match = MARKDOWN_HEADING_RE.match(line)

        if heading_match:
            hashes, title = heading_match.groups()

            # Do not allow arbitrary extracted headings to become the
            # root CV title.
            level = max(3, len(hashes))

            output.append(
                f"{'#' * level} {title.strip()}"
            )
            previous_blank = False
            continue

        # Clean weird bullet symbols while preserving Markdown bullets.
        if BULLET_RE.match(line):
            content = BULLET_RE.sub("", line).strip()

            if content:
                output.append(f"- {content}")

            previous_blank = False
            continue

        output.append(line)
        previous_blank = False

    result = "\n".join(output)

    result = re.sub(
        r"\n{3,}",
        "\n\n",
        result,
    )

    return result.strip()


def _ensure_cv_title(markdown: str) -> str:
    if not markdown:
        return ""

    # Prevent duplicate root title
    if re.match(
        r"^\s*#\s+CV\b",
        markdown,
        flags=re.IGNORECASE,
    ):
        return markdown.strip()

    return f"# CV\n\n{markdown.strip()}"


def _column_split_x(words: list[dict], page_width: float) -> float | None:
    """
    Find a vertical gutter wide enough to be a real column boundary
    (sidebar | main content), not just normal word spacing. Returns the
    split x-coordinate, or None if the page doesn't look multi-column.
    """
    if not words or page_width <= 0:
        return None

    spans = sorted((w["x0"], w["x1"]) for w in words)
    best_gap = 0.0
    best_split: float | None = None
    prev_end = spans[0][1]

    for x0, x1 in spans[1:]:
        gap = x0 - prev_end
        mid = (prev_end + x0) / 2
        # Only consider gaps roughly in the middle third of the page —
        # a gutter between a sidebar and the main column, not the margin.
        if gap > best_gap and page_width * 0.2 < mid < page_width * 0.8:
            best_gap = gap
            best_split = mid
        prev_end = max(prev_end, x1)

    if best_gap < page_width * 0.05:
        return None
    return best_split


def _words_to_lines(words: list[dict]) -> str:
    rows: dict[int, list[dict]] = defaultdict(list)
    for word in words:
        rows[round(word["top"] / 3)].append(word)
    lines = []
    for key in sorted(rows):
        row = sorted(rows[key], key=lambda w: w["x0"])
        lines.append(" ".join(w["text"] for w in row))
    return "\n".join(lines)


def _pdfplumber_to_markdown(data: bytes) -> str:
    """
    Column-aware fallback extraction.

    pymupdf4llm reads a page in a single left-to-right/top-to-bottom pass,
    which garbles real-world multi-column templates (sidebar + main
    content interleaved mid-sentence). This splits each page into columns
    by word position first, then reads each column top-to-bottom.
    """
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(keep_blank_chars=False)
            split_x = _column_split_x(words, page.width)
            if split_x is None:
                parts.append(page.extract_text() or "")
                continue
            left = [w for w in words if (w["x0"] + w["x1"]) / 2 < split_x]
            right = [w for w in words if (w["x0"] + w["x1"]) / 2 >= split_x]
            parts.append(_words_to_lines(left))
            parts.append(_words_to_lines(right))
    return _clean_text_artifacts("\n\n".join(p for p in parts if p.strip()))


def _parse_pdf(data: bytes) -> str:
    """
    Normal extraction first. Full OCR only when the result clearly
    appears corrupted. If content is still thin after that, try a
    column-aware fallback for multi-column/icon-heavy real-world layouts.
    """
    markdown = _pdf_to_markdown(
        data,
        force_ocr=False,
    )

    if _looks_corrupted(markdown):
        try:
            ocr_markdown = _pdf_to_markdown(
                data,
                force_ocr=True,
            )

            # Do not replace usable native extraction with empty /
            # obviously worse OCR.
            if (
                ocr_markdown.strip()
                and not _looks_corrupted(ocr_markdown)
            ):
                markdown = ocr_markdown

        except Exception:
            # OCR engine/language pack may not be installed.
            # Keep the native extraction instead of failing the parser.
            pass

    if len(markdown.strip()) < LOW_CONTENT_CHAR_THRESHOLD:
        try:
            layout_markdown = _pdfplumber_to_markdown(data)
        except Exception:
            layout_markdown = ""

        if len(layout_markdown.strip()) > len(markdown.strip()):
            markdown = layout_markdown

    return markdown


def _text_to_markdown(text: str) -> str:
    """
    Plain-text resume fallback.

    This is intentionally less aggressive than the old implementation:
    arbitrary lines are NOT automatically converted into bullets.
    """
    text = _clean_text_artifacts(text)

    lines = [
        line.strip()
        for line in text.splitlines()
    ]

    output: list[str] = []

    for line in lines:
        if not line:
            if output and output[-1] != "":
                output.append("")
            continue

        detected = _detect_section(line)

        if detected:
            section, remainder = detected
            if output and output[-1] != "":
                output.append("")

            output.append(f"## {section}")
            output.append("")
            if remainder:
                output.append(remainder)
            continue

        if BULLET_RE.match(line):
            content = BULLET_RE.sub("", line).strip()
            if content:
                output.append(f"- {content}")
            continue

        output.append(line)

    return "\n".join(output).strip()


def _docx_to_markdown(data: bytes) -> str:
    """
    Real DOCX extraction via python-docx (paragraphs + tables), instead of
    the previous silent UTF-8-decode-as-text fallback which mangled the
    zipped XML binary.
    """
    import io

    doc = _DocxDocument(io.BytesIO(data))
    output: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if output and output[-1] != "":
                output.append("")
            continue
        style = (para.style.name if para.style else "") or ""
        if style.casefold().startswith("heading") or style.casefold() == "title":
            output.append(f"### {text}")
        elif style.casefold().startswith(("list", "bullet")):
            output.append(f"- {text}")
        else:
            output.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                output.append("- " + " | ".join(cells))

    return _clean_text_artifacts("\n".join(output))


def _looks_like_docx(data: bytes, mime_type: str) -> bool:
    if mime_type.casefold() in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/docx",
    }:
        return True
    # DOCX is a zip archive; a bare zip magic number isn't proof enough,
    # so also check for the DOCX-specific internal part.
    return data[:4] == b"PK\x03\x04" and b"word/document.xml" in data


def _name_token(word: str) -> bool:
    token = word.strip(".,;:")
    if not token.isalpha():
        return False
    folded = token.casefold()
    if folded in _NAME_PARTICLES or (len(token) == 1):
        return True
    if token.isupper() and len(token) <= 15:
        return True
    return token[0].isupper() and token[1:].islower()


def _looks_like_person_name(line: str) -> bool:
    text = MARKDOWN_HEADING_RE.sub(r"\2", line).strip()
    if not text or text.casefold() in {"cv", "resume", "curriculum vitae"}:
        return False
    words = text.split()
    if not (2 <= len(words) <= 5):
        return False
    if any(word.casefold() in _JOB_WORDS for word in words):
        return False
    return all(_name_token(word) for word in words)


def _is_bare_name_continuation(line: str) -> bool:
    """A single bare word right after a redacted name line (a wrapped given
    name/surname split across two lines by the PDF layout)."""
    text = line.strip()
    if not text or text.startswith("#"):
        return False
    words = text.split()
    return len(words) == 1 and _name_token(words[0])


# Safety cap on the "## Contact" skip-scope in redact_pii(): if a later
# section heading in the CV isn't recognized (unusual wording, decorative
# icon regex misses, etc.), skip_contact would otherwise never turn back
# off and the entire rest of the CV silently gets dropped. A real contact
# block is a handful of lines, so stop unconditionally skipping well
# before that becomes catastrophic; the regex passes below
# (email/phone/url/dob) still strip real PII either way.
_CONTACT_SKIP_LINE_CAP = 15


def redact_pii(markdown: str) -> str:
    """Strip identifiers so LLM/embed never see the person. Ceiling: regex + header heuristics."""
    lines_out: list[str] = []
    skip_contact = False
    seen_section = False
    pending_name_continuation = False
    contact_lines_skipped = 0
    for raw in markdown.splitlines():
        heading = re.match(r"^##\s+(.+?)\s*$", raw)
        if heading:
            seen_section = True
            skip_contact = heading.group(1).strip().casefold() in _CONTACT_SECTIONS
            contact_lines_skipped = 0
            if skip_contact:
                pending_name_continuation = False
                continue
        elif skip_contact:
            contact_lines_skipped += 1
            if contact_lines_skipped > _CONTACT_SKIP_LINE_CAP:
                skip_contact = False
            else:
                continue
        if LABELED_PII_RE.match(raw):
            pending_name_continuation = False
            continue
        dropped_as_name = (not seen_section and _looks_like_person_name(raw)) or (
            raw.startswith("#") and _looks_like_person_name(raw)
        )
        if dropped_as_name:
            pending_name_continuation = True
            continue
        if pending_name_continuation and _is_bare_name_continuation(raw):
            pending_name_continuation = False
            continue
        pending_name_continuation = False
        lines_out.append(raw)

    text = "\n".join(lines_out)
    text = MARKDOWN_LINK_RE.sub("", text)
    text = EMAIL_RE.sub("", text)
    text = URL_RE.sub("", text)
    text = PHONE_RE.sub("", text)
    text = DOB_RE.sub("", text)
    cleaned: list[str] = []
    for line in text.splitlines():
        stripped = BULLET_RE.sub("", line).strip()
        if not stripped:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        cleaned.append(line.rstrip())
    return re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned)).strip()


def clean_markdown(markdown: str) -> str:
    return _ensure_cv_title(_normalize_markdown_sections(_clean_text_artifacts(markdown)))


def parse_resume_bytes(
    data: bytes,
    mime_type: str = "",
) -> dict:
    """PDF/DOCX/text bytes -> PII-stripped Markdown + parse-quality metadata.

    Skills are extracted from this markdown before summarize (see the
    ingest graph), not from the LLM-rewritten body.
    """
    if not data:
        return {"markdown": "", "metadata": {}}

    # Do NOT lstrip PDF bytes before parsing.
    # Whitespace before %PDF is unusual but stripping arbitrary binary
    # data is unnecessary.
    is_pdf = (
        mime_type.casefold() == "application/pdf"
        or data.lstrip()[:4] == b"%PDF"
    )

    if is_pdf:
        markdown = _parse_pdf(data)
    elif _looks_like_docx(data, mime_type):
        try:
            markdown = _docx_to_markdown(data)
        except Exception:
            # Corrupted/unsupported docx variant: fall back to best-effort
            # text decode rather than failing the whole ingest.
            markdown = _text_to_markdown(data.decode("utf-8", errors="replace"))
    else:
        text = data.decode(
            "utf-8",
            errors="replace",
        )
        markdown = _text_to_markdown(text)

    cleaned = clean_markdown(markdown)
    content_chars = len(cleaned.strip())
    metadata = {
        "content_chars": content_chars,
        "low_content": content_chars < LOW_CONTENT_CHAR_THRESHOLD,
    }

    return {"markdown": redact_pii(cleaned), "metadata": metadata}
